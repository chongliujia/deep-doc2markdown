import os
import uuid
import docx
from PIL import Image
import io

from config.config import IMAGES_DIR


class DocxExtractor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.text_content = []
        self.images = []
        self.structure = {}

    def extract_text(self):
        """Extract text from DOCX file with structure information."""
        doc = docx.Document(self.file_path)
        
        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                style_name = para.style.name
                self.text_content.append({
                    "index": i,
                    "content": para.text,
                    "style": style_name,
                    "type": "paragraph"
                })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell_idx, cell in enumerate(row.cells):
                    row_text.append(cell.text)
                table_text.append(row_text)
            
            if table_text:
                self.text_content.append({
                    "type": "table",
                    "content": table_text,
                    "index": table_idx
                })
        
        return self.text_content

    def extract_images(self):
        """Extract images from DOCX file."""
        doc = docx.Document(self.file_path)
        
        image_index = 0
        # Process all the inline shapes in the document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                
                # Generate a unique filename
                image_filename = f"docx_image_{uuid.uuid4()}.png"
                image_path = os.path.join(IMAGES_DIR, image_filename)
                
                # Open and save the image
                image = Image.open(io.BytesIO(image_data))
                image.save(image_path)
                
                # Store image metadata
                self.images.append({
                    "index": image_index,
                    "filename": image_filename,
                    "path": image_path,
                    "width": image.width,
                    "height": image.height
                })
                
                image_index += 1
        
        return self.images

    def extract_all(self):
        """Extract both text and images from DOCX."""
        self.extract_text()
        self.extract_images()
        
        return {
            "text": self.text_content,
            "images": self.images
        } 